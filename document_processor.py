import os
import hashlib
import pickle
from datetime import datetime, timedelta
from pathlib import Path
from typing import List, Union
import concurrent.futures
import logging

import fitz  # PyMuPDF
import pypdfium2
import camelot
import pdfplumber
from pdf2image import convert_from_path
import pytesseract
import docx  # python-docx for DOCX text extraction


# Configure logger
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentProcessor:
    def __init__(self,
                 cache_dir: Union[str, Path] = "./cache",
                 ocr_dpi: int = 200,
                 max_ocr_pages: int = 10,
                 enable_table_extraction: bool = True,
                 cache_expire_days: int = 7,
                 tesseract_cmd: str = None):
        self.cache_dir = Path(cache_dir)
        self.cache_dir.mkdir(parents=True, exist_ok=True)

        self.ocr_dpi = ocr_dpi
        self.max_ocr_pages = max_ocr_pages
        self.enable_table_extraction = enable_table_extraction
        self.cache_expire_days = cache_expire_days

        if tesseract_cmd:
            pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

    def process(self, files: List) -> List[str]:
        """
        Processes multiple files concurrently and returns list of extracted text chunks.
        """
        results = []
        max_workers = min(4, os.cpu_count() or 1)
        # Use ProcessPoolExecutor if any file is large (>10MB)
        executor_class = concurrent.futures.ThreadPoolExecutor
        if any(os.path.getsize(f.name) > 10 * 1024 * 1024 for f in files if hasattr(f, "name")):
            executor_class = concurrent.futures.ProcessPoolExecutor

        with executor_class(max_workers=max_workers) as executor:
            futures = {executor.submit(self._process_single_file, f): getattr(f, "name", repr(f)) for f in files}
            for future in concurrent.futures.as_completed(futures):
                fname = futures[future]
                try:
                    chunks = future.result()
                    results.extend(chunks)
                except Exception as e:
                    logger.error(f"Error processing file {fname}: {e}")

        logger.info(f"Total chunks extracted from all files: {len(results)}")
        return results

    def extract_text_from_file(self, filepath: str) -> str:
        """
        Extract text from a single file and return as a single string.
        """
        if not os.path.isfile(filepath):
            logger.error(f"Invalid file path: {filepath}")
            return ""

        logger.info(f"Extracting text from file {filepath}")
        
        ext = filepath.split(".")[-1].lower()
        text = ""
        
        if ext == "pdf":
            text = self._extract_pdf_text(filepath)
        elif ext == "docx":
            text = self._extract_docx_text(filepath)
        elif ext in ("txt", "md"):
            try:
                with open(filepath, "r", encoding="utf-8") as f:
                    text = f.read()
            except Exception as e:
                logger.error(f"Failed to read text file {filepath}: {e}")
                text = ""
        else:
            logger.warning(f"Unsupported file type {ext}: {filepath}")
            text = ""

        if not text.strip() and ext == "pdf":
            logger.info(f"No text found in {filepath}, trying OCR fallback")
            text = self._extract_pdf_ocr(filepath)

        if self.enable_table_extraction and ext == "pdf":
            text = self._append_tables_from_pdf(filepath, text)

        logger.info(f"Extracted {len(text)} characters from {filepath}")
        return text

    def _process_single_file(self, file) -> List[str]:
        """
        Process a single file: try cache, extract text (PDF, DOCX, TXT, MD), parse tables,
        OCR fallback, then chunk and cache results.
        """
        filename = getattr(file, "name", None)
        if not filename or not os.path.isfile(filename):
            logger.error(f"Invalid file object or path: {file}")
            return []

        logger.info(f"Processing file {filename}")
        with open(filename, "rb") as f:
            content = f.read()

        file_hash = self._generate_hash(content)
        cache_path = self.cache_dir / f"{file_hash}.pkl"
        if self._is_cache_valid(cache_path):
            logger.info(f"Loading cached data for {filename}")
            return self._load_cache(cache_path)

        ext = filename.split(".")[-1].lower()
        text = ""
        if ext == "pdf":
            text = self._extract_pdf_text(filename)
        elif ext == "docx":
            text = self._extract_docx_text(filename)
        elif ext in ("txt", "md"):
            try:
                with fitz.open(filename, "r", encoding="utf-8") as f:
                    text = f.read()
            except Exception as e:
                logger.error(f"Failed to read text file {filename}: {e}")
                text = ""
        else:
            logger.warning(f"Unsupported file type {ext}: {filename}")
            text = ""

        if not text.strip() and ext == "pdf":
            logger.info(f"No text found in {filename}, trying OCR fallback")
            text = self._extract_pdf_ocr(filename)

        if self.enable_table_extraction and ext == "pdf":
            text = self._append_tables_from_pdf(filename, text)

        # Simple chunking by paragraphs
        chunks = [chunk.strip() for chunk in text.split("\n\n") if chunk.strip()]

        self._save_cache(chunks, cache_path)

        logger.info(f"Extracted {len(chunks)} chunks from {filename}")
        return chunks

    def _extract_pdf_text(self, filepath: str) -> str:
        """
        Try extract PDF text with pypdfium2 then fallback to pdfplumber then fitz text extraction.
        """
        print(f"🔍 [TEXT EXTRACTION] Starting text extraction for: {filepath}")
        
        # pypdfium2
        print(f"📄 [TEXT EXTRACTION] Trying pypdfium2 extraction...")
        try:
            pdf = pypdfium2.PdfDocument(filepath)
            try:
                pages_text = [page.get_textpage().get_text_range() for page in pdf]
                combined = "\n".join(pages_text).strip()
                if combined:
                    print(f"✅ [TEXT EXTRACTION] pypdfium2 successful - extracted {len(combined)} characters")
                    print(f"📝 [TEXT EXTRACTION] First 200 chars: {combined[:200]}...")
                    return combined
                else:
                    print(f"⚠️ [TEXT EXTRACTION] pypdfium2 returned empty text")
            finally:
                pdf.close()
        except Exception as e:
            print(f"❌ [TEXT EXTRACTION] pypdfium2 failed: {e}")

        # pdfplumber
        print(f"📄 [TEXT EXTRACTION] Trying pdfplumber extraction...")
        try:
            texts = []
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        texts.append(page_text)
            text = "\n".join(texts).strip()
            if text:
                print(f"✅ [TEXT EXTRACTION] pdfplumber successful - extracted {len(text)} characters")
                print(f"📝 [TEXT EXTRACTION] First 200 chars: {text[:200]}...")
                return text
            else:
                print(f"⚠️ [TEXT EXTRACTION] pdfplumber returned empty text")
        except Exception as e:
            print(f"❌ [TEXT EXTRACTION] pdfplumber failed: {e}")

        # PyMuPDF (fitz)
        print(f"📄 [TEXT EXTRACTION] Trying PyMuPDF (fitz) extraction...")
        try:
            with fitz.open(filepath) as doc:
                texts = [page.get_text("text") for page in doc]
            result = "\n".join(texts).strip()
            print(f"✅ [TEXT EXTRACTION] PyMuPDF successful - extracted {len(result)} characters")
            print(f"📝 [TEXT EXTRACTION] First 200 chars: {result[:200]}...")
            return result
        except Exception as e:
            print(f"❌ [TEXT EXTRACTION] PyMuPDF failed: {e}")
            return ""

    def _extract_pdf_ocr(self, filepath: str) -> str:
        """
        OCR extracting PDF pages (limited to max_ocr_pages) using pdf2image and pytesseract with concurrency.
        """
        try:
            with fitz.open(filepath) as doc:
                num_pages = len(doc)
                if num_pages > self.max_ocr_pages:
                    logger.warning(f"Skipping OCR for {filepath} with {num_pages} pages > max OCR pages {self.max_ocr_pages}")
                    return ""

            images = convert_from_path(filepath, dpi=self.ocr_dpi,
                                       first_page=1, last_page=min(num_pages, self.max_ocr_pages))
            text_chunks = []
            with concurrent.futures.ThreadPoolExecutor(max_workers=2) as executor:
                futures = [executor.submit(pytesseract.image_to_string, img) for img in images]
                for future in concurrent.futures.as_completed(futures):
                    text_chunks.append(future.result())

            return "\n".join(text_chunks).strip()
        except Exception as e:
            logger.error(f"OCR extraction failed for {filepath}: {e}")
            return ""

    def _append_tables_from_pdf(self, filepath: str, base_text: str) -> str:
        """
        Extract tables from PDF using camelot (stream + lattice) or pdfplumber fallback,
        append markdown tables to base text.
        """
        tables_md = []

        # Try camelot with "stream"
        try:
            tables = camelot.read_pdf(filepath, flavor="stream", pages="all")
            for table in tables:
                md = table.df.to_markdown(index=False)
                if md:
                    tables_md.append(md)
        except Exception as e:
            logger.debug(f"Camelot stream extraction failed: {e}")
            # Try camelot with "lattice"
            try:
                tables = camelot.read_pdf(filepath, flavor="lattice", pages="all")
                for table in tables:
                    md = table.df.to_markdown(index=False)
                    if md:
                        tables_md.append(md)
            except Exception as e:
                logger.debug(f"Camelot lattice extraction failed: {e}")
                # Fallback to pdfplumber
                try:
                    with pdfplumber.open(filepath) as pdf:
                        for page in pdf.pages:
                            for table in page.extract_tables():
                                if table and len(table) > 1:
                                    df = table_to_df(table)
                                    md = df.to_markdown(index=False)
                                    tables_md.append(md)
                except Exception as e:
                    logger.debug(f"pdfplumber table extraction failed for {filepath}: {e}")

        if tables_md:
            return f"{base_text}\n\n" + "\n\n".join(tables_md)
        else:
            return base_text

    def _extract_docx_text(self, filepath: str) -> str:
        """
        Extract clean text from DOCX file using python-docx.
        """
        try:
            doc = docx.Document(filepath)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs).strip()
        except Exception as e:
            logger.error(f"DOCX extraction failed for {filepath}: {e}")
            return ""

    @staticmethod
    def _generate_hash(content: bytes) -> str:
        return hashlib.sha256(content).hexdigest()

    def _save_cache(self, chunks: List[str], cache_path: Path):
        try:
            with open(cache_path, "wb") as f:
                pickle.dump({
                    "timestamp": datetime.now().timestamp(),
                    "chunks": chunks
                }, f)
        except Exception as e:
            logger.error(f"Failed to save cache {cache_path}: {e}")

    def _load_cache(self, cache_path: Path) -> List[str]:
        try:
            with open(cache_path, "rb") as f:
                data = pickle.load(f)
                # Support both dict and list for backward compatibility
                if isinstance(data, dict):
                    return data.get("chunks", [])
                elif isinstance(data, list):
                    return data
                else:
                    return []
        except Exception as e:
            logger.error(f"Failed to load cache {cache_path}: {e}")
            return []


    def _is_cache_valid(self, cache_path: Path) -> bool:
        if not cache_path.exists():
            return False
        age = datetime.now() - datetime.fromtimestamp(cache_path.stat().st_mtime)
        return age < timedelta(days=self.cache_expire_days)


def table_to_df(table):
    """
    Convert extracted table (list of lists) to pandas DataFrame with header.
    """
    import pandas as pd

    if not table or len(table) < 2:
        return pd.DataFrame()

    header, *rows = table
    return pd.DataFrame(rows, columns=header)


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python document_processor.py <file1> [file2 ...]")
        exit(1)

    processor = DocumentProcessor()
    file_paths = sys.argv[1:]
    class FileMock:
        """Simple wrapper to mimic file object with .name attribute."""
        def __init__(self, path):
            self.name = path

    files = [FileMock(fp) for fp in file_paths]
    chunks = processor.process(files)

    print(f"\nExtracted {len(chunks)} text chunks:")
    for i, chunk in enumerate(chunks[:5], 1):
        print(f"\n--- Chunk {i} ---\n{chunk}")

